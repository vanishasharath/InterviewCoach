# backend/extractor.py

import re
import io
import logging
from typing import Dict, List, Optional, Set
from pathlib import Path

import spacy
from PyPDF2 import PdfReader
from docx import Document

# -------------------------------------------------------------------
# Logging
# -------------------------------------------------------------------
logger = logging.getLogger(__name__)

# -------------------------------------------------------------------
# Configuration
# -------------------------------------------------------------------
SKILL_CORPUS_PATH = "backend/skill_corpus.txt"
DEFAULT_SPACY_MODEL = "en_core_web_sm"
MIN_TEXT_LENGTH = 50

# -------------------------------------------------------------------
# Load spaCy model
# -------------------------------------------------------------------
try:
    nlp = spacy.load(DEFAULT_SPACY_MODEL)
    logger.info(f"Loaded spaCy model: {DEFAULT_SPACY_MODEL}")
except OSError as e:
    logger.error(
        f"spaCy model '{DEFAULT_SPACY_MODEL}' not found. "
        f"Install it using: python -m spacy download {DEFAULT_SPACY_MODEL}"
    )
    raise e

# -------------------------------------------------------------------
# Skill Corpus Loader
# -------------------------------------------------------------------
class SkillCorpusLoader:
    _corpus_cache: Optional[Set[str]] = None

    @classmethod
    def load_corpus(cls, filepath: str = SKILL_CORPUS_PATH) -> Set[str]:
        if cls._corpus_cache is not None:
            return cls._corpus_cache

        corpus_path = Path(filepath)
        if not corpus_path.exists():
            logger.warning("Skill corpus file not found. Using empty corpus.")
            cls._corpus_cache = set()
            return cls._corpus_cache

        with open(corpus_path, "r", encoding="utf-8") as f:
            cls._corpus_cache = {
                line.strip().lower()
                for line in f
                if line.strip() and not line.startswith("#")
            }

        logger.info(f"Loaded {len(cls._corpus_cache)} skills from corpus")
        return cls._corpus_cache


SKILL_CORPUS = SkillCorpusLoader.load_corpus()

# -------------------------------------------------------------------
# File Extraction
# -------------------------------------------------------------------
def extract_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from PDF, DOCX, or TXT files.
    """
    if not file_bytes or not filename:
        raise ValueError("Empty file or filename")

    filename = filename.lower()

    if filename.endswith(".pdf"):
        text = _extract_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        text = _extract_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        text = _extract_from_text(file_bytes)
    elif filename.endswith(".doc"):
        raise ValueError("Legacy .doc format not supported. Please upload .docx or PDF.")
    else:
        raise ValueError("Unsupported file format")

    text = text.strip()

    if not text or len(text) < MIN_TEXT_LENGTH:
        raise ValueError(
            "Failed to extract meaningful text from resume. "
            "Please ensure the file is not empty or corrupted."
        )

    return text.lower()


def _extract_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    reader = PdfReader(io.BytesIO(file_bytes))

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def _extract_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def _extract_from_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")

# -------------------------------------------------------------------
# NLP Utilities
# -------------------------------------------------------------------
def normalize_text(text: str) -> str:
    text = re.sub(r"[^a-z0-9+.\s]", " ", text.lower())
    return re.sub(r"\s+", " ", text).strip()


def extract_skills(text: str, corpus: Optional[Set[str]] = None) -> List[str]:
    if not text:
        return []

    corpus = corpus or SKILL_CORPUS
    if not corpus:
        return []

    doc = nlp(text)
    found_skills: Set[str] = set()

    for chunk in doc.noun_chunks:
        phrase = normalize_text(chunk.text)
        if phrase in corpus:
            found_skills.add(phrase)
        for word in phrase.split():
            if word in corpus:
                found_skills.add(word)

    for token in doc:
        if not token.is_stop and not token.is_punct:
            tok = normalize_text(token.text)
            if tok in corpus:
                found_skills.add(tok)

    normalized_text = normalize_text(text)
    for skill in corpus:
        if len(skill.split()) > 1:
            if re.search(rf"\b{re.escape(skill)}\b", normalized_text):
                found_skills.add(skill)

    return sorted(found_skills)

# -------------------------------------------------------------------
# Information Extraction
# -------------------------------------------------------------------
def extract_experience_years(text: str) -> float:
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)",
        r"over\s+(\d+(?:\.\d+)?)\s*(?:years?|yrs?)",
        r"more than\s+(\d+(?:\.\d+)?)\s*(?:years?|yrs?)",
    ]

    for pattern in patterns:
        for match in re.finditer(pattern, text.lower()):
            try:
                years = float(match.group(1))
                if 0 < years <= 50:
                    return years
            except ValueError:
                pass

    return 0.0


def extract_education(text: str) -> List[str]:
    keywords = [
        "bachelor", "master", "phd", "b.tech", "m.tech",
        "b.e", "m.e", "b.sc", "m.sc", "mba"
    ]
    found = set()
    text_lower = text.lower()

    for key in keywords:
        if key in text_lower:
            found.add(key)

    return list(found)


def extract_contact_info(text: str) -> Dict[str, Optional[str]]:
    info = {"email": None, "phone": None, "linkedin": None}

    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if email:
        info["email"] = email.group()

    phone = re.search(r"\+?\d[\d\s\-()]{8,}\d", text)
    if phone:
        info["phone"] = phone.group()

    linkedin = re.search(r"linkedin\.com/in/[\w-]+", text.lower())
    if linkedin:
        info["linkedin"] = linkedin.group()

    return info


def categorize_skills(skills: List[str], text: str) -> Dict[str, List[str]]:
    text_lower = text.lower()
    project_kw = ["project", "developed", "built", "implemented"]
    cert_kw = ["certified", "certification"]

    result = {
        "resume_skills": skills,
        "project_skills": [],
        "certification_skills": [],
    }

    for skill in skills:
        if any(k in text_lower for k in project_kw):
            result["project_skills"].append(skill)
        if any(k in text_lower for k in cert_kw):
            result["certification_skills"].append(skill)

    return result


def extract_information(text: str) -> Dict:
    if not text:
        return {
            "experience_years": 0.0,
            "skills": {},
            "education": [],
            "contact_info": {},
            "total_skills_found": 0,
        }

    all_skills = extract_skills(text)
    categorized = categorize_skills(all_skills, text)

    return {
        "experience_years": extract_experience_years(text),
        "skills": categorized,
        "education": extract_education(text),
        "contact_info": extract_contact_info(text),
        "total_skills_found": len(all_skills),
    }
