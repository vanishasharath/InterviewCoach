# backend/extractor.py

import re
import io
import logging
from typing import Dict, List, Optional, Set
from pathlib import Path
import spacy

# Set up logging
logger = logging.getLogger(__name__)

# Configuration
SKILL_CORPUS_PATH = "backend/skill_corpus.txt"
DEFAULT_SPACY_MODEL = "en_core_web_lg"

# Load spaCy model with error handling
try:
    nlp = spacy.load(DEFAULT_SPACY_MODEL)
    logger.info(f"Loaded spaCy model: {DEFAULT_SPACY_MODEL}")
except OSError:
    logger.error(f"spaCy model '{DEFAULT_SPACY_MODEL}' not found. Please install it with: python -m spacy download {DEFAULT_SPACY_MODEL}")
    raise


class SkillCorpusLoader:
    """Load and manage skill corpus"""
    
    _corpus_cache: Optional[Set[str]] = None
    
    @classmethod
    def load_corpus(cls, filepath: str = SKILL_CORPUS_PATH) -> Set[str]:
        """
        Load skill corpus from file with caching
        
        Args:
            filepath: Path to skill corpus file
            
        Returns:
            Set of skills in lowercase
        """
        if cls._corpus_cache is not None:
            return cls._corpus_cache
        
        try:
            corpus_path = Path(filepath)
            if not corpus_path.exists():
                logger.warning(f"Skill corpus file not found at {filepath}. Using empty corpus.")
                cls._corpus_cache = set()
                return cls._corpus_cache
            
            with open(corpus_path, "r", encoding="utf-8") as f:
                cls._corpus_cache = {
                    line.strip().lower()
                    for line in f.readlines()
                    if line.strip() and not line.strip().startswith("#")
                }
            
            logger.info(f"Loaded {len(cls._corpus_cache)} skills from corpus")
            return cls._corpus_cache
            
        except Exception as e:
            logger.error(f"Error loading skill corpus: {e}")
            cls._corpus_cache = set()
            return cls._corpus_cache
    
    @classmethod
    def reload_corpus(cls, filepath: str = SKILL_CORPUS_PATH) -> Set[str]:
        """Force reload the corpus"""
        cls._corpus_cache = None
        return cls.load_corpus(filepath)


# Load universal skill corpus
SKILL_CORPUS = SkillCorpusLoader.load_corpus()


def extract_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from various file formats
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename with extension
        
    Returns:
        Extracted text content in lowercase
    """
    if not file_bytes:
        logger.warning("Empty file bytes provided")
        return ""
    
    if not filename:
        logger.warning("No filename provided")
        return ""
    
    text = ""
    filename_lower = filename.lower()
    
    try:
        if filename_lower.endswith(".pdf"):
            text = _extract_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            text = _extract_from_docx(file_bytes)
        elif filename_lower.endswith((".txt", ".text")):
            text = _extract_from_text(file_bytes)
        elif filename_lower.endswith(".doc"):
            logger.warning("Legacy .doc format detected. Consider converting to .docx")
            text = _extract_from_text(file_bytes)  # Fallback
        else:
            logger.warning(f"Unsupported file format: {filename}")
            text = _extract_from_text(file_bytes)  # Try as plain text
            
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}", exc_info=True)
        return ""
    
    # Clean and normalize text
    text = text.strip()
    if not text:
        logger.warning(f"No text extracted from {filename}")
    
    return text.lower()


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
        raise
    
    text_parts = []
    try:
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        
        if not pdf_reader.pages:
            logger.warning("PDF has no pages")
            return ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Error extracting text from PDF page {page_num + 1}: {e}")
                continue
        
        logger.info(f"Extracted text from {len(text_parts)} PDF pages")
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        raise


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX"""
    try:
        import docx
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        raise
    
    text_parts = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        logger.info(f"Extracted text from DOCX with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise


def _extract_from_text(file_bytes: bytes) -> str:
    """Extract text from plain text file"""
    try:
        # Try UTF-8 first
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decoding failed, trying with latin-1")
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            logger.warning(f"Latin-1 decoding failed: {e}, using ignore errors")
            return file_bytes.decode("utf-8", errors="ignore")


def normalize_text(text: str) -> str:
    """
    Normalize text for skill matching
    
    Args:
        text: Input text
        
    Returns:
        Normalized text
    """
    # Remove special characters but keep +, ., and spaces
    normalized = re.sub(r"[^a-z0-9+.\s]", " ", text.lower())
    # Collapse multiple spaces
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def extract_skills(text: str, corpus: Optional[Set[str]] = None) -> List[str]:
    """
    Use spaCy to extract skills from text based on skill corpus
    
    Args:
        text: Input text to analyze
        corpus: Optional skill corpus (uses global SKILL_CORPUS if not provided)
        
    Returns:
        Sorted list of identified skills
    """
    if not text or not text.strip():
        logger.warning("Empty text provided for skill extraction")
        return []
    
    if corpus is None:
        corpus = SKILL_CORPUS
    
    if not corpus:
        logger.warning("Empty skill corpus - no skills will be extracted")
        return []
    
    try:
        doc = nlp(text)
        found_skills: Set[str] = set()
        
        # Extract skills from noun chunks
        for chunk in doc.noun_chunks:
            phrase = normalize_text(chunk.text)
            if phrase in corpus:
                found_skills.add(phrase)
            
            # Also check individual words in the chunk
            for word in phrase.split():
                if len(word) > 1 and word in corpus:
                    found_skills.add(word)
        
        # Extract skills from individual tokens
        for token in doc:
            # Skip stop words and punctuation
            if token.is_stop or token.is_punct:
                continue
            
            normalized_token = normalize_text(token.text)
            if normalized_token and normalized_token in corpus:
                found_skills.add(normalized_token)
        
        # Extract skills from named entities (ORG, PRODUCT, etc.)
        for ent in doc.ents:
            if ent.label_ in ["ORG", "PRODUCT", "GPE"]:
                normalized_entity = normalize_text(ent.text)
                if normalized_entity in corpus:
                    found_skills.add(normalized_entity)
        
        # Check for multi-word skills directly in text
        text_normalized = normalize_text(text)
        for skill in corpus:
            if len(skill.split()) > 1:  # Multi-word skills
                if re.search(r'\b' + re.escape(skill) + r'\b', text_normalized):
                    found_skills.add(skill)
        
        logger.debug(f"Extracted {len(found_skills)} skills from text")
        return sorted(found_skills)
        
    except Exception as e:
        logger.error(f"Error extracting skills: {e}", exc_info=True)
        return []


def extract_experience_years(text: str) -> float:
    """
    Extract numeric years of experience from resume or job description
    
    Args:
        text: Input text
        
    Returns:
        Number of years of experience (0.0 if not found)
    """
    if not text:
        return 0.0
    
    # Comprehensive patterns for experience extraction
    patterns = [
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)",
        r"experience[:\-\s]*(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)?",
        r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s+(?:in|of|with)",
        r"over\s+(\d+(?:\.\d+)?)\s+(?:years?|yrs?)",
        r"more than\s+(\d+(?:\.\d+)?)\s+(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\s*\+\s*(?:years?|yrs?)",
        r"(\d+(?:\.\d+)?)\s*to\s*\d+\s*(?:years?|yrs?)",  # Range: take first number
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text.lower())
        for match in matches:
            try:
                years = float(match.group(1))
                # Sanity check: reasonable experience range
                if 0 < years <= 50:
                    logger.debug(f"Extracted experience: {years} years")
                    return years
            except (ValueError, IndexError):
                continue
    
    logger.debug("No experience years found in text")
    return 0.0


def extract_education(text: str) -> List[str]:
    """
    Extract education information from text
    
    Args:
        text: Input text
        
    Returns:
        List of education qualifications
    """
    education_keywords = [
        r"bachelor(?:'s)?(?:\s+of)?",
        r"master(?:'s)?(?:\s+of)?",
        r"phd|ph\.d|doctorate",
        r"associate(?:'s)?(?:\s+degree)?",
        r"mba",
        r"b\.?s\.?|m\.?s\.?|b\.?a\.?|m\.?a\.?",
        r"b\.?tech|m\.?tech|b\.?e\.?|m\.?e\.?",
    ]
    
    education = []
    text_lower = text.lower()
    
    for keyword in education_keywords:
        matches = re.finditer(keyword, text_lower)
        for match in matches:
            # Extract surrounding context (up to 50 chars)
            start = max(0, match.start() - 10)
            end = min(len(text_lower), match.end() + 40)
            context = text[start:end].strip()
            education.append(context)
    
    return list(set(education))  # Remove duplicates


def extract_contact_info(text: str) -> Dict[str, Optional[str]]:
    """
    Extract contact information from text
    
    Args:
        text: Input text
        
    Returns:
        Dictionary with email, phone, and linkedin
    """
    contact_info = {
        "email": None,
        "phone": None,
        "linkedin": None
    }
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact_info["email"] = email_match.group(0)
    
    # Phone pattern (various formats)
    phone_patterns = [
        r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
        r'\b\(\d{3}\)\s*\d{3}[-.]?\d{4}\b',
        r'\+\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4}\b'
    ]
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group(0)
            break
    
    # LinkedIn pattern
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin_match = re.search(linkedin_pattern, text.lower())
    if linkedin_match:
        contact_info["linkedin"] = linkedin_match.group(0)
    
    return contact_info


def categorize_skills(skills: List[str], text: str) -> Dict[str, List[str]]:
    """
    Categorize skills based on context in the text
    
    Args:
        skills: List of all extracted skills
        text: Original text for context
        
    Returns:
        Dictionary with categorized skills
    """
    text_lower = text.lower()
    
    # Keywords for different categories
    project_keywords = ["project", "developed", "built", "created", "implemented", "designed"]
    cert_keywords = ["certified", "certification", "certificate", "licensed"]
    
    categorized = {
        "resume_skills": list(skills),  # All skills
        "project_skills": [],
        "certification_skills": []
    }
    
    # Find skills mentioned near project keywords
    for skill in skills:
        skill_pattern = r'\b' + re.escape(skill) + r'\b'
        for match in re.finditer(skill_pattern, text_lower):
            start = max(0, match.start() - 100)
            end = min(len(text_lower), match.end() + 100)
            context = text_lower[start:end]
            
            # Check if in project context
            if any(keyword in context for keyword in project_keywords):
                if skill not in categorized["project_skills"]:
                    categorized["project_skills"].append(skill)
            
            # Check if in certification context
            if any(keyword in context for keyword in cert_keywords):
                if skill not in categorized["certification_skills"]:
                    categorized["certification_skills"].append(skill)
    
    return categorized


def extract_information(text: str) -> Dict:
    """
    Extract comprehensive information from resume text
    
    Args:
        text: Resume text content
        
    Returns:
        Dictionary containing extracted information
    """
    if not text or not text.strip():
        logger.warning("Empty text provided for information extraction")
        return {
            "experience_years": 0.0,
            "skills": {
                "resume_skills": [],
                "project_skills": [],
                "certification_skills": [],
            },
            "education": [],
            "contact_info": {
                "email": None,
                "phone": None,
                "linkedin": None
            },
            "error": "Empty text provided"
        }
    
    try:
        # Extract all skills
        all_skills = extract_skills(text)
        
        # Categorize skills based on context
        categorized_skills = categorize_skills(all_skills, text)
        
        # Extract other information
        experience_years = extract_experience_years(text)
        education = extract_education(text)
        contact_info = extract_contact_info(text)
        
        result = {
            "experience_years": experience_years,
            "skills": categorized_skills,
            "education": education,
            "contact_info": contact_info,
            "total_skills_found": len(all_skills)
        }
        
        logger.info(f"Successfully extracted information: {len(all_skills)} skills, {experience_years} years experience")
        return result
        
    except Exception as e:
        logger.error(f"Error extracting information: {e}", exc_info=True)
        return {
            "experience_years": 0.0,
            "skills": {
                "resume_skills": [],
                "project_skills": [],
                "certification_skills": [],
            },
            "education": [],
            "contact_info": {
                "email": None,
                "phone": None,
                "linkedin": None
            },
            "error": str(e)
        }


# Example usage and testing
if __name__ == "__main__":
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Test with sample text
    sample_text = """
    John Doe
    Email: john.doe@example.com
    Phone: (555) 123-4567
    LinkedIn: linkedin.com/in/johndoe
    
    Senior Software Engineer with 7+ years of experience in Python and web development.
    
    Education:
    - Master's in Computer Science
    - Bachelor's in Software Engineering
    
    Skills: Python, Django, Flask, JavaScript, React, PostgreSQL, Docker, AWS, Git
    
    Projects:
    - Developed a web application using Django and React
    - Built microservices with Docker and Kubernetes
    
    Certifications:
    - AWS Certified Solutions Architect
    - Certified Kubernetes Administrator
    """
    
    print("\n=== Testing Extraction ===")
    result = extract_information(sample_text.lower())
    
    print(f"\nExperience: {result['experience_years']} years")
    print(f"\nTotal Skills Found: {result.get('total_skills_found', 0)}")
    print(f"\nAll Skills: {', '.join(result['skills']['resume_skills'][:10])}")
    print(f"\nProject Skills: {', '.join(result['skills']['project_skills'][:5])}")
    print(f"\nCertification Skills: {', '.join(result['skills']['certification_skills'][:5])}")
    print(f"\nEducation: {result.get('education', [])[:2]}")
    print(f"\nContact: {result.get('contact_info', {})}")